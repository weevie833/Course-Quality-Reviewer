"""
Course Quality Reviewer — FastAPI backend
"""

import json
import os
import re
import shutil
from contextlib import asynccontextmanager
from datetime import datetime
from pathlib import Path
from typing import Any

import anthropic
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.responses import FileResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).parent

_env_file = BASE_DIR / ".env"
if _env_file.exists():
    for _line in _env_file.read_text().splitlines():
        if "=" in _line and not _line.startswith("#"):
            _k, _v = _line.split("=", 1)
            _key, _val = _k.strip(), _v.strip()
            if _val and not os.environ.get(_key):
                os.environ[_key] = _val

MODEL = os.environ.get("ANTHROPIC_MODEL", "claude-sonnet-4-6")

EXTRACTED_COURSES_DIR = BASE_DIR / "extracted_courses"
EXTRACTED_COURSES_DIR.mkdir(exist_ok=True)
MAX_COURSE_VERSIONS = 5

client = anthropic.Anthropic()

# ---------------------------------------------------------------------------
# App state
# ---------------------------------------------------------------------------

app_state: dict[str, Any] = {}


@asynccontextmanager
async def lifespan(app: FastAPI):
    yield


app = FastAPI(lifespan=lifespan)
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")


# ---------------------------------------------------------------------------
# Pydantic models
# ---------------------------------------------------------------------------

class Message(BaseModel):
    role: str
    content: str


class ChatRequest(BaseModel):
    message: str
    history: list[Message] = []


class EvaluateRequest(BaseModel):
    course_id: str
    timestamp: str


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/")
async def index():
    return FileResponse(BASE_DIR / "static" / "index.html")


@app.get("/health")
async def health():
    return {"status": "ok", "model": MODEL}


_EASTER_EGG_RE = re.compile(r"^\s*about the cqr\s*[.!?]?\s*$", re.IGNORECASE)
_EASTER_EGG_TEXT = """\
The CQR was created by Steve Covello, Senior Instructional Designer at UNH-CPSO in May, 2026. \
It was developed locally on a MacBook Pro using Claude Code.

* Co-orientation principles were derived from decades of research in information systems design \
by Dr. Brenda Dervin (Sense-Making Methodology). User-based design principles were derived from \
Dr. Michael Nilan's research at Syracuse University's iSchool.
* Service orientation was created by Steve Covello from 30+ years in creative client relations.
* AI integration principles were contributed from the "Controlling AI in Instruction" e-book \
authored by Steve Covello (2025).
* Rich media pedagogy principles from the "Teaching with Rich Media" e-book by Steve Covello (2019).

The CQR was tested by the Instructional Design team at UNH-CPSO and overseen by \
Reta Chaffee, Director of Educational Technology."""

_CHAT_SYSTEM_PROMPT = """\
You are the Course Quality Reviewer assistant at the University of New Hampshire College of \
Professional Studies. You help course designers, instructional designers, and faculty understand \
evaluation results, ask questions about course content, and improve course quality.

Answer directly and concisely from the course content provided. If a question cannot be answered \
from the available content, say so clearly rather than speculating.\
"""


@app.post("/chat")
async def chat(req: ChatRequest):
    if not req.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty.")

    if _EASTER_EGG_RE.match(req.message):
        def _egg():
            yield _EASTER_EGG_TEXT
        return StreamingResponse(_egg(), media_type="text/plain")

    # Build system prompt — inject current course content if available
    system = _CHAT_SYSTEM_PROMPT
    current = app_state.get("current_course")
    if current:
        version_dir = (
            EXTRACTED_COURSES_DIR / current["course_id"] / current["timestamp"]
        )
        if version_dir.exists():
            items = _load_course_content(version_dir)
            content_text = _format_content_for_prompt(items, current["course_id"])
            system += (
                f"\n\nCURRENT COURSE: {current['course_id']} — {current['course_title']}\n\n"
                f"COURSE CONTENT:\n\n{content_text}"
            )

    messages = [{"role": m.role, "content": m.content} for m in req.history]
    messages = messages[-6:] if len(messages) > 6 else messages
    messages.append({"role": "user", "content": req.message})

    def generate():
        try:
            with client.messages.stream(
                model=MODEL,
                max_tokens=4096,
                system=system,
                messages=messages,
            ) as stream:
                for text in stream.text_stream:
                    yield text
        except anthropic.RateLimitError:
            yield "The system is temporarily rate-limited. Please wait a moment and try again."
        except anthropic.APIError as e:
            yield f"API error: {getattr(e, 'message', str(e))}"
        except Exception as e:
            yield f"Unexpected error: {type(e).__name__}: {e}"

    return StreamingResponse(generate(), media_type="text/plain")


# ---------------------------------------------------------------------------
# CQR — content loading helpers
# ---------------------------------------------------------------------------

def _strip_html(html_text: str) -> str:
    from html.parser import HTMLParser

    class _Stripper(HTMLParser):
        def __init__(self):
            super().__init__(convert_charrefs=True)
            self._parts: list[str] = []
            self._skip = False

        def handle_starttag(self, tag, attrs):
            if tag in ("script", "style"):
                self._skip = True
            elif tag == "img":
                attr_map = dict(attrs)
                alt = attr_map.get("alt", "").strip()
                role = attr_map.get("role", "")
                if role == "presentation" or alt == "":
                    self._parts.append("[Image: no alt text]")
                else:
                    self._parts.append(f"[Image: alt='{alt}']")
            elif tag == "iframe":
                attr_map = dict(attrs)
                title = attr_map.get("title", "").strip()
                src = attr_map.get("src", "").lower()
                if "youtube" in src:
                    kind = "YouTube video"
                elif "kaltura" in src or "brightcove" in src or "vimeo" in src:
                    kind = "Kaltura/hosted video"
                elif "$canvas_course_reference$" in src or "external_tools" in src:
                    kind = "Canvas-embedded video (likely Kaltura)"
                else:
                    kind = "embedded video"
                label = f'"{title}"' if title else "(untitled)"
                self._parts.append(f"[{kind}: {label} — captioning requires human verification]")

        def handle_endtag(self, tag):
            if tag in ("script", "style"):
                self._skip = False

        def handle_data(self, data):
            if not self._skip:
                self._parts.append(data)

        def get_text(self):
            return re.sub(r'\s+', ' ', "".join(self._parts)).strip()

    stripper = _Stripper()
    stripper.feed(html_text)
    return stripper.get_text()


def _parse_file_meta(raw: str) -> tuple[str, str, str]:
    def _get(key: str) -> str:
        m = re.search(rf'<!--\s*{re.escape(key)}:\s*(.+?)\s*-->', raw, re.IGNORECASE)
        return m.group(1).strip() if m else ""
    return _get("CONTENT-TYPE"), _get("TITLE"), _get("MODULE")


def _module_sort_key(item: dict) -> tuple:
    """Sort key: Syllabus → Course Overview → Module 1-N → everything else."""
    ct = item["content_type"]
    mod = item["module"]
    if ct == "Syllabus":
        return (0, 0, "")
    if mod == "Course Overview & Resources":
        return (1, 0, "")
    m = re.match(r"Module\s+(\d+)", mod, re.IGNORECASE)
    if m:
        return (2, int(m.group(1)), item["title"])
    return (3, 0, mod + item["title"])


def _load_course_content(version_dir: Path, char_limit_per_item: int = 2000) -> list[dict]:
    items = []
    for html_file in sorted(version_dir.glob("*.html")):
        raw = html_file.read_text(errors="replace")
        content_type, title, module = _parse_file_meta(raw)
        if not content_type:
            content_type = "Unknown"
        if not title:
            title = html_file.stem
        text = _strip_html(raw)
        if content_type != "Syllabus" and len(text) > char_limit_per_item:
            truncated = text[:char_limit_per_item]
            last_space = truncated.rfind(' ')
            text = truncated[:last_space] + ' [...]'
        if text:
            items.append({"content_type": content_type, "title": title, "module": module, "text": text})
    items.sort(key=_module_sort_key)
    return items


def _format_content_for_prompt(items: list[dict], course_id: str) -> str:
    if not items:
        return "(No course content found)"
    lines = []
    for item in items:
        module_label = f" — {item['module']}" if item['module'] else ""
        lines.append(f"=== {item['content_type']}: {item['title']}{module_label} ===")
        lines.append(item['text'])
        lines.append("")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# CQR — evaluation
# ---------------------------------------------------------------------------

from quality_standards import STANDARDS as CQR_STANDARDS

_EVAL_SYSTEM_PROMPT = """\
You are conducting a structured course quality evaluation for the Course Quality Reviewer (CQR) \
at the University of New Hampshire. Your role is to locate and report evidence from the extracted \
course content that relates to each quality sub-standard.

Rules:
- For each sub-standard, write the evidence as a bulleted list — one bullet per distinct piece of \
evidence found. If nothing was found, use a single bullet: "- None found."
- Every bullet must describe the location hierarchically before naming what was found. \
Use this format depending on content type:
    Module content: "In the [Module Name], in the '[Content Title]' [page/assignment/discussion], [what was found]."
    Syllabus: "In the [Section Name] section of the syllabus, [paragraph or table], [what was found]."
  Examples:
    "In the Course Overview & Resources module, in the 'Course Resources' page, links to rubrics and APA resources are consolidated."
    "In the Expectations and Conduct section of the syllabus, the Code of Conduct paragraph references the Student Rights, Rules, and Responsibilities."
- Do not make recommendations. Do not suggest improvements. Report evidence only.
- Be specific: quote phrases, cite exact titles, and name all location levels.
- Evaluate from the perspective of a first-time student visiting the course.
- Do NOT include the verdict inside the evidence field.
- Verdict values:
  - "Met": the course clearly and sufficiently addresses the sub-standard.
  - "Partially Met": the course addresses the sub-standard but incompletely or with notable gaps.
  - "Not Met": evidence is absent or insufficient to satisfy the sub-standard.
- When the verdict is "Partially Met" or "Not Met", populate the "gaps" field with a bulleted list \
of what is absent or insufficient. Each gap bullet begins with "- Missing: " or "- Insufficient: ". \
When the verdict is "Met", set "gaps" to "".
- The verdict is a provisional signal to guide the reviewer's attention — not a final determination.
- Output ONLY valid JSON. No markdown fences, no prose outside the JSON.

Output format (one object per sub-standard in order):
{
  "substandards": [
    {
      "id": "1.1",
      "description": "...",
      "evidence": "- In Module 1, in the 'Tasks for getting started' page, ...\\n- In the syllabus, ...",
      "gaps": "",
      "verdict": "Met"
    },
    ...
  ]
}
"""


def _run_standard_evaluation(
    standard_num: int,
    course_id: str,
    course_title: str,
    content_text: str,
) -> dict:
    std = CQR_STANDARDS[standard_num]
    substandards_block = "\n\n".join(
        f'Sub-standard {ss["id"]}: {ss["description"]}\n'
        f'Evaluator guidance: {ss["interpretation"]}'
        for ss in std["substandards"]
    )
    user_message = (
        f"COURSE: {course_id} — {course_title}\n\n"
        f"SUB-STANDARDS TO EVALUATE (Standard {standard_num}):\n\n"
        f"{substandards_block}\n\n"
        f"COURSE CONTENT:\n\n{content_text}"
    )
    response = client.messages.create(
        model=MODEL,
        max_tokens=8192,
        system=_EVAL_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": user_message}],
    )
    raw = response.content[0].text.strip()
    # Strip markdown fences if Claude wraps the JSON despite instructions
    if raw.startswith("```"):
        raw = re.sub(r'^```[a-z]*\n?', '', raw)
        raw = re.sub(r'\n?```$', '', raw).strip()
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        # Attempt repair for minor issues (unescaped quotes, trailing commas, etc.)
        try:
            from json_repair import repair_json
            return json.loads(repair_json(raw))
        except Exception:
            import logging
            logging.error("CQR eval JSON unrecoverable. stop_reason=%s raw=%s", response.stop_reason, raw[-500:])
            raise json.JSONDecodeError("Unrecoverable JSON from model", raw, 0)


def _read_canvas_course_info(imscc_bytes: bytes) -> tuple[str, str]:
    import io
    import re as _re
    import zipfile as _zf
    import xml.etree.ElementTree as _ET

    def _clean_course_id(raw_code: str) -> str:
        m = _re.search(r"([A-Za-z]+)\s+(\d{3,4})", raw_code)
        return f"{m.group(1).upper()}-{m.group(2)}" if m else raw_code

    try:
        with _zf.ZipFile(io.BytesIO(imscc_bytes)) as zf:
            try:
                xml_data = zf.read("course_settings/course_settings.xml")
                root = _ET.fromstring(xml_data)
                ns = "http://canvas.instructure.com/xsd/cccv1p0"

                def _find(tag):
                    r = root.find(f"{{{ns}}}{tag}")
                    return r if r is not None else root.find(tag)

                code_el  = _find("course_code")
                title_el = _find("title")
                raw_code = (code_el.text  or "").strip() if code_el  is not None else ""
                title    = (title_el.text or "").strip() if title_el is not None else ""
                if raw_code:
                    return _clean_course_id(raw_code), title
            except (KeyError, _ET.ParseError):
                pass

            try:
                xml_data = zf.read("imsmanifest.xml")
                root = _ET.fromstring(xml_data)
                for el in root.iter():
                    if el.tag.split("}")[-1] == "string" and el.text and el.text.strip():
                        t = el.text.strip()
                        return _clean_course_id(t), t
            except (KeyError, _ET.ParseError):
                pass
    except Exception:
        pass
    return "", ""


# ---------------------------------------------------------------------------
# CQR — upload, session, evaluation endpoints
# ---------------------------------------------------------------------------

@app.post("/upload")
async def upload_course(
    file: UploadFile = File(...),
    discard_oldest: str = Form("false"),
):
    if not (file.filename or "").lower().endswith(".imscc"):
        raise HTTPException(status_code=400, detail="Only .imscc files are accepted.")

    import asyncio
    import tempfile
    from ingest_imscc import ingest_imscc

    content = await file.read()
    canvas_code, canvas_title = _read_canvas_course_info(content)

    with tempfile.NamedTemporaryFile(suffix=".imscc", delete=False) as tmp:
        tmp_path = Path(tmp.name)
        tmp.write(content)

    course_id    = canvas_code or "UNKNOWN"
    course_title = canvas_title or course_id
    course_dir   = EXTRACTED_COURSES_DIR / course_id
    course_dir.mkdir(exist_ok=True)

    existing = sorted([d for d in course_dir.iterdir() if d.is_dir()])

    if len(existing) >= MAX_COURSE_VERSIONS and discard_oldest.lower() != "true":
        tmp_path.unlink(missing_ok=True)
        return {
            "at_version_limit": True,
            "version_count": len(existing),
            "course_id": course_id,
            "course_title": course_title,
        }

    if len(existing) >= MAX_COURSE_VERSIONS and discard_oldest.lower() == "true":
        shutil.rmtree(existing[0])
        existing = existing[1:]

    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    version_dir = course_dir / ts
    version_dir.mkdir()

    try:
        result = await asyncio.to_thread(ingest_imscc, tmp_path, course_id, version_dir)
    except (FileNotFoundError, ValueError) as exc:
        raise HTTPException(status_code=422, detail=str(exc))
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))
    finally:
        tmp_path.unlink(missing_ok=True)

    course_title = canvas_title or result.get("course_title", course_id)

    meta = {
        "course_id": course_id,
        "course_title": course_title,
        "timestamp": ts,
        "incomplete_standards": [],
        "evaluation": {},
    }
    (version_dir / "meta.json").write_text(json.dumps(meta))

    prior_ts = existing[-1].name if existing else None
    app_state["current_course"] = {
        "course_id": course_id,
        "course_title": course_title,
        "timestamp": ts,
        "prior_version_ts": prior_ts,
        "version_count": len(existing) + 1,
        "incomplete_standards": [],
        "evaluation": {},
        "at_version_limit": False,
    }
    return app_state["current_course"]


@app.get("/course/current")
async def get_current_course():
    return app_state.get("current_course") or {}


STANDARD_TITLES = {
    1: "Course Overview and Introduction",
    2: "Learning Objectives",
    3: "Assessment and Measurement",
    4: "Instructional Materials",
    5: "Learning Activities and Learner Interaction",
    6: "Course Technology",
    7: "Learner Support",
    8: "Usability",
}

TEMPLATE_PATH = BASE_DIR / "templates" / "UNH_CQR_Template.docx"


class DocxRequest(BaseModel):
    course_id: str
    timestamp: str


def _build_docx(course_title: str, evaluation: dict) -> bytes:
    import io
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.oxml import OxmlElement

    doc = Document(str(TEMPLATE_PATH))

    # The "[Course ID and title]" line is a body paragraph styled "Header" (not a page header)
    # Replace the placeholder run text in-place before clearing the rest of the body
    for p in doc.paragraphs:
        for run in p.runs:
            if "[Course ID and title]" in run.text:
                run.text = run.text.replace("[Course ID and title]", course_title)

    # Remove all body paragraphs from the first bold sub-standard line onwards
    body = doc.element.body
    paragraphs = doc.paragraphs
    cut_from = None
    for i, p in enumerate(paragraphs):
        if p.runs and any(r.bold for r in p.runs) and re.match(r"\d+\.\d+:", p.text.strip()):
            cut_from = i
            break
    if cut_from is not None:
        for p in paragraphs[cut_from:]:
            body.remove(p._element)

    def add_para(text="", bold=False, size_pt=None, color=None, indent_pt=None, bullet=False):
        p = doc.add_paragraph()
        if bullet:
            # Manual bullet with hanging indent
            pf = p.paragraph_format
            from docx.shared import Inches
            pf.left_indent = Inches(0.25)
            pf.first_line_indent = Inches(-0.25)
            run = p.add_run(f"• {text}")
        else:
            run = p.add_run(text)
        if bold:
            run.bold = True
        if size_pt:
            run.font.size = Pt(size_pt)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if indent_pt and not bullet:
            from docx.shared import Pt as _Pt
            p.paragraph_format.left_indent = _Pt(indent_pt)
        return p

    def split_bullets(text):
        return [l.strip().lstrip("-• ").strip() for l in (text or "").split("\n") if l.strip()]

    for std_num in range(1, 9):
        std_data = evaluation.get(str(std_num))
        if not std_data:
            continue

        add_para()
        add_para(f"Standard {std_num}: {STANDARD_TITLES[std_num]}", bold=True, size_pt=12)

        for ss in std_data.get("substandards", []):
            add_para()
            add_para(f"{ss['id']}: {ss['description']}", bold=True)

            # Verdict: label bold, value explicitly non-bold
            p = doc.add_paragraph()
            r1 = p.add_run("Verdict: ")
            r1.bold = True
            r2 = p.add_run(ss.get("verdict", ""))
            r2.bold = False

            for line in split_bullets(ss.get("evidence", "")):
                add_para(line, bullet=True)

            gaps = ss.get("gaps", "")
            if gaps and gaps.strip():
                add_para()
                add_para("Gaps", bold=True, color=(0xB4, 0x53, 0x09))
                for line in split_bullets(gaps):
                    p = doc.add_paragraph()
                    from docx.shared import Inches
                    pf = p.paragraph_format
                    pf.left_indent = Inches(0.25)
                    pf.first_line_indent = Inches(-0.25)
                    run = p.add_run(f"• {line}")
                    run.font.color.rgb = RGBColor(0x92, 0x40, 0x0E)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.post("/download/docx")
async def download_docx(req: DocxRequest):
    version_dir = EXTRACTED_COURSES_DIR / req.course_id / req.timestamp
    meta_path = version_dir / "meta.json"
    if not meta_path.exists():
        raise HTTPException(status_code=404, detail="No evaluation found for this course version.")

    meta = json.loads(meta_path.read_text())
    course_title = meta.get("course_title", req.course_id)
    evaluation = meta.get("evaluation", {})

    import asyncio as _asyncio
    docx_bytes = await _asyncio.to_thread(_build_docx, course_title, evaluation)

    safe_title = re.sub(r"[^\w\s-]", "", course_title)[:60].strip()
    filename = f"CQR_{safe_title}.docx"
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/evaluate/{standard_num}")
async def evaluate_standard(standard_num: int, req: EvaluateRequest):
    if standard_num < 1 or standard_num > 8:
        raise HTTPException(status_code=400, detail="Standard number must be 1–8.")

    version_dir = EXTRACTED_COURSES_DIR / req.course_id / req.timestamp
    if not version_dir.exists():
        raise HTTPException(status_code=404, detail=f"Course version not found: {req.course_id}/{req.timestamp}")

    items = _load_course_content(version_dir)
    content_text = _format_content_for_prompt(items, req.course_id)

    meta_path = version_dir / "meta.json"
    course_title = req.course_id
    if meta_path.exists():
        meta = json.loads(meta_path.read_text())
        course_title = meta.get("course_title", req.course_id)

    import asyncio as _asyncio
    try:
        result = await _asyncio.to_thread(
            _run_standard_evaluation,
            standard_num, req.course_id, course_title, content_text,
        )
    except json.JSONDecodeError as exc:
        raise HTTPException(status_code=502, detail=f"Model returned invalid JSON: {exc}")
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc))

    if meta_path.exists():
        meta = json.loads(meta_path.read_text())
        meta["evaluation"][str(standard_num)] = result
        meta_path.write_text(json.dumps(meta))
        if app_state.get("current_course"):
            app_state["current_course"]["evaluation"] = meta["evaluation"]

    return {"standard": standard_num, "result": result}
